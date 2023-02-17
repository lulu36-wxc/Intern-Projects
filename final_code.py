#-*- codeing = utf-8 -*-
from bs4 import BeautifulSoup
import re
import urllib.request
import urllib.error
import xlwt
import sqlite3
import datetime
import time
import os
import requests
import pdfplumber
from docx import Document
from openpyxl import Workbook
import shutil
import xlrd
import pandas as pd
from xlrd import xldate_as_datetime
from xlrd import xldate_as_tuple
from selenium import webdriver
import numpy as np
from win32com import client as wc
#pip install python-docx


last_time = input("Please input the last time you searched (yyyymmdd):")    #输入上次查询日期，程序仅返回该日期后的数据
today = datetime.date.today().strftime('%Y%m%d')                            #返回今日日期

#定义Cookie、User-Agent、Accept
head = {
        "Cookie": "wzws_cid=bea9a57941ad6181d9630672efea0f63fe0ea35e4e3bf3153e4068706686365a94fde1ba0d602eec8bc085c9ecdb54d77acf98e1660111f7f3dea1939a242c5dd064ed67920b57501428ac06bd88a4ed",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.159 Safari/537.36 Edg/92.0.902.78"
    }

#正则表达式
findLink = re.compile(r'<a href="(.*?)">')
findCHN = re.compile(u'[\u4e00-\u9fa5]')

#新建dataframe，将所有支行链接里的地名拼音与支行名称一一对应
d = {
    'place': ['shanghai','wuhan','chongqing','haerbin','changsha','guiyang','wulumuqi',
              'xiamen','tianjin','guangzhou','shijiazhuang','hangzhou','nanchang','lasa',
              'shenzhen','shenyang','chengdu','taiyuan','fuzhou','nanning','lanzhou','dalian',
              'nanjing','xian','huhehaote','hefei','haikou','xining','qingdao','jinan',
              'beijing','changchun','zhengzhou','kunming','yinchuan','ningbo'],
    'name': ['上海分行','武汉分行','重庆营业管理部','哈尔滨中心支行','长沙中心支行','贵阳中心支行','乌鲁木齐中心支行',
             '厦门市中心支行','天津分行','广州分行','石家庄中心支行','杭州中心支行','南昌中心支行','拉萨中心支行',
             '深圳市中心支行','沈阳分行','成都分行','太原中心支行','福州中心支行','南宁中心支行','兰州中心支行',
             '大连市中心支行','南京分行','西安分行','呼和浩特中心支行','合肥中心支行','海口中心支行','西宁中心支行',
             '青岛市中心支行','济南分行','北京营业管理部','长春中心支行','郑州中心支行','昆明中心支行','银川中心支行',
             '宁波市中心支行']
}
df_inf = pd.DataFrame(d)

#定义筛选反洗钱相关关键词
key_words = ['客户身份','大额','可疑交易','身份不明','反洗钱','保密','泄露','交易信息',
            '真实性','完整性','可追溯性','商户资料','商户信息','不明身份','交易监测','异常交易','匿名账户',
             '假名账户','未按规定处理异议','风险等级']


information = open('information.txt','r',encoding='utf-8')      #读取information文件中已写好的路径
outfile = open('Remark.txt','w')                                #所以特殊提醒将写入Remark文件
path_1 = information.read().split('\n')[0].replace(' ','')      #定义路径

#判断运行文件夹里是否存在以今日日期、今日日期+(Excel)、Need Check命名的文件夹，
#若已存在则删除新建文件夹，若不存在则新建文件夹
if os.path.exists(path_1+today):
    shutil.rmtree(path_1+today)
    shutil.rmtree(path_1+today+'(Excel)')
    os.mkdir(path_1+today)
    os.mkdir(path_1+today+'(Excel)')
else:
    os.mkdir(path_1+today)
    os.mkdir(path_1+today+'(Excel)')

if os.path.exists(path_1+'Need Check\\'):
    shutil.rmtree(path_1+'Need Check\\')
    os.mkdir(path_1+'Need Check\\')
else:
    os.mkdir(path_1+'Need Check\\')

def main():
    #读取各个支行的链接地址
    infile = open(path_1+"links.txt", 'r',encoding='utf-8')                 
    data = infile.read().split("\n")
    infile.close()
    n_link = 0
    #循环链接地址，将所有支行数据爬出
    for i in data:
        baseurl = i
        cookie_1 = getData(baseurl)
        if cookie_1 == 'invalid':
            restart(data,n_link)
            break
        n_link += 1
        
    outfile.close()
    combine_data(today)     #合并所有数据并筛选

def restart(data,n_link):
    new_cookie = input('Please input the new Cookie:')
    head['Cookie'] = new_cookie
    n = n_link
    for i in range(n,len(data)):
        cookie_1 = getData(data[i])
        if cookie_1 == 'invalid':
            restart(data,n_link)
        n_link += 1

#提取网页源代码
def askURL(url):
    request = urllib.request.Request(url,headers = head)
    html = ""
    try:
        response = urllib.request.urlopen(request)
        html = response.read().decode('utf-8')
    except urllib.error.URLError as e:
        if hasattr(e,"code"):
            print(e.code)
        if hasattr(e,"reason"):
            print(e.reason)

    return html

#爬取数据
def getData(baseurl):
    for x in range(1,100):
        print(x)
        place = baseurl.split('/')[2].split('.')[0]   #提取该支行链接所在地的拼音
        if place == 'xiamen':
            url = baseurl
        else:
            url = baseurl.replace('index1.html','')+'index'+str(x)+".html"
        html_1 = askURL(url)        #提取网页源代码
        soup_1 = BeautifulSoup(html_1,"html.parser")
        count_1 = 0         #新增行政处罚链接的个数
        count_2 = 0         #已抓取行政处罚链接的个数
        index = 0
        date_time_list = []

        if soup_1.find_all('font', class_ = "hei12") == []:
            cookie_1 = 'invalid'
            return cookie_1
        
        #查找该支行链接里在上次查询后新增行政处罚链接的个数，最终个数为count_1个
        for n in soup_1.find_all('td', class_ = "hei12jj"):
            date_time = re.findall(r'\d\d\d\d-\d\d-\d\d',str(n))
            if date_time == []:
                continue
            date_int = date_time[0].replace('-', '')
            if date_int > last_time:
                count_1 += 1        #若行政处罚链接对应日期晚于上次查询日期则count_1加一
                date_time_list.append(date_int)
                
        #定位该支行链接里所有存在行政处罚链接的区域并一一循环
        for item_1 in soup_1.find_all('font', class_ = "hei12"):
            count_2 += 1            #记录在该支行链接里已抓取行政处罚链接的个数
            if count_2 > count_1:   #若已抓取行政处罚链接个数大于新增行政处罚链接个数则结束循环，停止循环该支行链接并进入下一个支行链接
                return
            item_1 = str(item_1)
            
            inf_1 = re.findall(findLink,item_1)[0]        #查找该区域的行政处罚链接
            link_1 = "http://"+place+".pbc.gov.cn/" + inf_1.split('\"')[0]  #定义行政处罚链接
            
            html_2 = askURL(link_1)         #提取该行政处罚链接的网页源代码
            soup_2 = BeautifulSoup(html_2, "html.parser")

            #定位该行政处罚链接里所有存在文件链接的区域
            for item_2 in soup_2.find_all('td', class_ = "hei14jj"):
                item_2 = str(item_2)
                inf_2 = re.findall(findLink,item_2)     #查找该区域里的对应文件链接

                if count_2 != 1:
                    if date_time_list[count_2-1] == date_time_list[count_2-2]:
                        index += 1
                        name = place+date_time_list[count_2-1]+'('+str(index/2).replace('.0','')+')'
                    else:
                        index = 0
                        name = place+date_time_list[count_2-1]
                else:
                    index = 0
                    name = place+date_time_list[count_2-1]
                #若该区域不存在文件链接，则该网页为一张表格
                if inf_2 == []:
                    file_type = 'html'      #文件类型为html
                    download_html(link_1, name, file_type)  #下载该网页的html文件

                    #对七个网页为表格的地区进行判断并导入Excel文件
                    if place == 'dalian':
                        dalian(name)
                    if place == 'qingdao':
                        qingdao(name)
                    if place == 'wulumuqi':
                        wulumuqi(name)
                    if place == 'kunming':
                        kunming(name)
                    if place == 'nanjing':
                        nanjing(name)
                    if place == 'nanchang':
                        nanchang(name)
                    if place == 'beijing':
                        beijing(name)
                    modify_excel_html(name,link_1)      #对Excel文件进行格式修改并放入以今日日期+(Excel)命名的文件夹里
                    os.remove(path_1+today+'\\'+name+'.xlsx')
                    print(link_1)

                #若该区域存在文件链接，则循环全部文件链接
                elif inf_2 != []:
                    for i in inf_2:
                        link_2 = "http://"+place+".pbc.gov.cn/" + i.split('\"')[0]      #定义文件链接
                        file_type = link_2.split('/')[-1].split('.')[1]     #定义下载文件类型(pdf,doc,docx,wps,xls,xlsx,et)
                        download_file(link_2,name,file_type)        #下载文件
                        if file_type.lower() == 'pdf':
                            pdf_to_excel(name,link_1)       #将pdf文件转换成Excel
                        if file_type.lower() == 'docx': 
                            docx_to_excel(name,link_1)      #将docx文件转换成Excel
                        if file_type.lower() == 'xls' or file_type.lower() == 'xlsx' or file_type.lower() == 'et':
                            modify_excel(name,file_type,link_1)     #修改xls，xlsx，et文件的格式以方便合并所有文件               
                        if file_type.lower() == 'doc' or file_type.lower() == 'wps':
                            doc_to_docx(name)               #将doc文件转换成docx
                            docx_to_excel(name,link_1)      #将docx文件转换成Excel
                            os.remove(path_1+today+'\\'+name+'.doc')
                        
                        print(link_2)
                index += 1


#下载文件
def download_file(url,name,file_type):
    r = requests.get(url, headers = head)
    #若文件类型为wps，则下载成doc文件
    if file_type.lower() == 'wps':
        with open(path_1+today+'\\'+name+'.doc', 'wb') as f:
            f.write(r.content)
    #若文件类型不为wps，则直接下载文件
    else:
        with open(path_1+today+'\\'+name+'.'+file_type, 'wb') as f:
            f.write(r.content)

#下载html文件
def download_html(url,name,file_type):
    r = askURL(url)
    with open(path_1+today+'\\'+name+'.'+file_type, 'wb') as f:
        f.write(bytes(r,encoding='utf-8'))

#将pdf转换成Excel文件
def pdf_to_excel(name,url):
    workbook = xlwt.Workbook()             #新建Excel文件
    sheet = workbook.add_sheet('Sheet1')   #新建Excel文件中的工作表并命名为'Sheet1'
    path = path_1+today+'\\'+name+'.pdf'   #定义pdf文件的路径
    pdf = pdfplumber.open(path)            #读取pdf文件
    
    i = 0               #记录在写入Excel文件时的行数
    nrow = 0            #定义pdf文件里表格的行数
    ncol = 0            #定义pdf文件里表格的列数
    full_table = []     #记录pdf文件里表格的内容
    
    #循环pdf文件里的表格并写入Excel文件
    for page in pdf.pages:                   
        for table in page.extract_tables():
            for row in table:
                nrow += 1
                ncol = len(row)
                n_null = 0        #定义表格里在第row行值为空的个数
                full_table += row
                for j in range(len(row)):
                    if row[j] == '':    #若第row行第j列的值为空则n_null加一
                        n_null += 1
                    if n_null == ncol:  #若第row行里值为空的个数等于表格列数，即第row行为空时，nrow减一
                        nrow -= 1
                    sheet.write(i, j, row[j])   #写入Excel文件
                i += 1
    if full_table == []:    #若表格为空，则该pdf为扫描件需要人工查看
        print('\nWarning: '+name+'.pdf is a scanned file. Please check it manually, thank you.\n')
        outfile.write(name+'.pdf is a scanned file. Please check it manually, thank you.\n\n')
        pdf.close()
        if os.path.exists(path_1+'Need Check\\'+name+'.pdf'):
            return
        shutil.move(path, path_1+'Need Check\\')    #将该pdf移入Need Check文件夹以方便后续人工查看
    else:
        sheet.write(0,ncol,'参考来源')      #若表格不为空，则添加"参考来源"列并写入该pdf文件所属的行政处罚链接
        for i in range(1,nrow):
            sheet.write(i,ncol,url)
        pdf.close()
        workbook.save(path_1+today+'(Excel)\\'+name+'.xls')     #保存Excel文件

#将docx转换成Excel文件
def docx_to_excel(name,url):
    workbook = xlwt.Workbook()              #新建Excel文件
    sheet = workbook.add_sheet('Sheet1')    #新建Excel文件中的工作表并命名为'Sheet1'
    path = path_1+today+'\\'+name+'.docx'   #定义docx文件的路径
    document = Document(path)               #读取docx文件
    
    nrow = 0        #定义docx文件里表格的行数
    i = 0           #记录在写入Excel文件时的行数
    x = 0           #定义变量x以判断该docx文件是否需要调用docx_text_excel函数进行格式修改
    a = 0           #定义变量a以判断该docx文件是否存在"备注"列
    total = len(document.tables)    #定义docx文件里表格的总行数
    
    for index in range(0, total):
        for row in document.tables[index].rows:
            nrow += 1
            ncol = len(row.cells)   #定义docx文件里表格的列数
            j = 0                   #记录在写入Excel文件时的列数
            n_null = 0              #定义表格里在第row行值为空的个数
            for grid in row.cells:
                if grid.text == '':
                    n_null += 1     #若第row行第j列的值为空则n_null加一
                if n_null == ncol:
                    nrow -= 1       #若第row行里值为空的个数等于表格列数，即第row行为空时，nrow减一
                if grid.text.replace(' ','').replace('\n','') == '备注':
                    a = 1           #若第row行第j列的值为"备注"，则将变量a赋值为1，即代表表格里存在"备注"列
                    
                #若表格里第一行存在两个值为空，或者表格里第一行存在一个值为空并且该表格无"备注"列，
                #则该docx遗失"处罚决定日期"相关数据，需要调用docx_text_excel函数进行格式修改，将变量x赋值为1
                if (i == 1 and n_null == 2) or (i == 1 and n_null == 1 and a == 0):
                    x = 1
                sheet.write(i,j,grid.text)  #写入Excel文件
                j += 1
            i += 1
    if ncol < 6:    #若表格列数小于6，则该docx文件格式有问题，需要人工查看
        print('\nWarning: the format of '+name+'.docx has some problem. Please check it manually, thank you.\n')
        outfile.write('The format of '+name+'.docx has some problem. Please check it manually, thank you.\n\n')
        if os.path.exists(path_1+'Need Check\\'+name+'.docx'):
            return
        shutil.move(path, path_1+'Need Check\\')    #将该docx移入Need Check文件夹以方便后续人工查看
    else:
        sheet.write(0,ncol,'参考来源')      #若表格不为空，则添加"参考来源"列并写入该docx文件所属的行政处罚链接
        for i in range(1,nrow):
            sheet.write(i,ncol,url)
        workbook.save(path_1+today+'(Excel)\\'+name+'.xls')     #保存Excel文件

        #若x的值为1，即该docx遗失"处罚决定日期"相关数据，则调用docx_text_excel函数进行格式修改
        if x == 1:
            docx_text_excel(name)

#若docx文件遗失"处罚决定日期"相关数据，则调用该函数进行修改
def docx_text_excel(name):
    df = pd.read_excel(path_1+today+'(Excel)\\'+name+'.xls')    #读取该docx已转换的Excel文件
    w2 = wc.Dispatch('Word.Application')
    doc2 = w2.Documents.Open(path_1+today+'\\'+name+'.docx')
    doc2.SaveAs(path_1+today+'\\'+name+'.txt', 4)               #将该docx文件转换成text文件
    doc2.Close()
    
    with open(path_1+today+'\\'+name+'.txt', 'r') as f:
        datalist=[]             #定义text文件里的数据
        data = f.readlines()
        for i in data:
            datalist.append(i.strip('\n'))
        time = []               #定义text文件里与"处罚决定日期"相关的数据
        for content in datalist:    #提取该text文件里与"处罚决定日期"相关的数据
            content = content.replace("年", "/").replace("月", "/").replace("日", "").replace("-", "/").replace(".","/")
            date = re.findall(re.compile('(\d{4}/\d{1,2}/\d{1,2})'), content)
            if date != []:
                time.append(date[0])
    if '备注' in df.columns:        #若已转换的Excel文件里存在"备注"列，则将"处罚决定日期"相关数据写入倒数第三列
        for i in range(len(time)):
            df[df.columns[-3]] = time[i]
    else:       #若已转换的Excel文件里不存在"备注"列，则将"处罚决定日期"相关数据写入倒数第二列
        for i in range(len(time)):
            df[df.columns[-2]] = time[i]

    df.to_excel(path_1+today+'(Excel)\\'+name+'.xls',index=False)   #导出格式正确的Excel
    os.remove(path_1+today+'\\'+name+'.txt')

#将doc转换成docx文件
def doc_to_docx(name):
    w = wc.Dispatch('Word.Application')
    doc = w.Documents.Open(path_1+today+'\\'+name+'.doc')
    doc.SaveAs(path_1+today+'\\'+name+'.docx',16)
    doc.Close()

#修改xls，xlsx，et文件的格式
def modify_excel(name,file_type,url):
    workbook = xlwt.Workbook()              #新建Excel文件
    sheet = workbook.add_sheet('Sheet1')    #新建Excel文件中的工作表并命名为'Sheet1'
    df = pd.read_excel(path_1+today+'\\'+name+'.'+file_type)    #读取Excel文件
    
    nrows = df.shape[0]       #定义Excel文件里表格的总行数
    ncols = df.columns.size   #定义Excel文件里表格的列数
    i = 0                     #记录在写入Excel文件时的行数
    n = 0                     #定义Excel文件里有效内容的行数

    for row in range(nrows):
        x = str(df.iloc[row,0]).replace('\n','')            #提取表格里第一列的内容
        #若第row行第一列的值不为"序号"或"当事人名称"相关并有效内容行数n为0，则第row行为表格标题或其他无效内容，进入下一行
        if x != '序号' and '当事人名称' not in x and n == 0:
            continue
        n += 1
        n_null = 0      #定义表格里在第row行值为空的个数
        for j in range(ncols):
            content = str(df.iloc[row,j]).replace(' ','')
            content = content.replace('\n','')      #定义第row行第j列的值
            if '填表' in content:     #若第row行第j列的值存在"填表"，则重新赋值为''
                content = ''
            if content == '' or content.lower() == 'nan':
                n_null += 1           #若第row行第j列的值为空则n_null加一
            if n_null == ncols:
                n -= 1                #若第row行里值为空的个数等于表格列数，即第row行为空时，n减一
            sheet.write(i, j, content)  #写入Excel文件
        i += 1

    sheet.write(0,ncols,'参考来源')     #添加"参考来源"列并写入该Excel文件所属的行政处罚链接
    for i in range(1,n):
        sheet.write(i,ncols,url)
    workbook.save(path_1+today+'(Excel)\\'+name+'.xls')     #保存文件

#将以下各个支行的html文件转换为Excel文件
#大连
def dalian(name):
    url=path_1+today+'\\'+name+'.html'
    l=pd.read_html(url) #用pandas安装包读取html文件，读取成功之后会在一个list里面
    df=l[9]  #在list中筛选出相对应的内容
    count=0 
    df=df.drop(df[[7]],axis=1) #删除备注列
    df=df.dropna(axis=0,how="all") #删除所有内容为空的行
    for i in range(len(df)):
        if df[[0]].iloc[i][0].isdigit() != True:
            count+=1
    df=df.iloc[count:,1:] #选取序号列为数字的行，即表格里的具体内容，不读取表头和序号列
    df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关', '处罚决定日期'] #重新对整个表格的列命名
    df.to_excel(path_1+today+'\\'+name+'.xlsx') #写入对应的Excel文件


#青岛
def qingdao(name):
    url=path_1+today+'\\'+name+'.html'
    l=pd.read_html(url) #用pandas安装包读取html文件，读取成功之后会在一个list里面
    df=l[9] #在list中筛选出相对应的内容
    count=0
    df=df.drop(df[[7,8]],axis=1) #删掉联系电话和备注列
    df=df.dropna(axis=0,how="all")  #删除所有内容为空的行
    for i in range(len(df)):
        if df[[0]].iloc[i][0].isdigit() != True:
            count+=1
    df=df.iloc[count:,1:] #选取序号列为数字的行，即表格里的具体内容，不读取表头和序号列
    df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关', '处罚决定日期'] #重新对整个表格的列命名
    df.to_excel(path_1+today+'\\'+name+'.xlsx') #写入对应的Excel文件
 

#乌鲁木齐
def wulumuqi(name):
    url=path_1+today+'\\'+name+'.html'
    l=pd.read_html(url) #用pandas安装包读取html文件，读取成功之后会在一个list里面
    df=l[9] #在list中筛选出相对应的内容
    count=0
    df=df.drop(df[[7]],axis=1) #删掉备注列
    df=df.dropna(axis=0,how="any") #删除所有内容为空的行
    df=df.iloc[1:,1:] # 读取表格具体内容，不读取表头和序号列
    df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关', '处罚决定日期'] #重新对整个表格的列命名
    df.to_excel(path_1+today+'\\'+name+'.xlsx') #写入对应的Excel文件


#昆明
def kunming(name):
    url=path_1+today+'\\'+name+'.html'
    l=pd.read_html(url) #用pandas安装包读取html文件，读取成功之后会在一个list里面
    df=l[9] #在list中筛选出相对应的内容
    count=0
    if len(df.columns)==6: #针对表格格式为六列的情况
        df=df.dropna(axis=0,how="any") #删掉任意列为空的行
        for i in range(len(df)):
            if df[[0]].iloc[i][0].isdigit() != True: 
                count+=1
        df=df.iloc[1:] #选取具体表格内容，不读取掉序号列
        df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关', '处罚决定日期'] #重新对整个表格的列命名
    elif len(df.columns)==9: #针对表格格式为九列的情况
        df=df.drop(df[[7,8]],axis=1) #去掉是否公示列和备注列
        df=df.dropna(axis=0,how="any") #删掉任意内容为空的行
        for i in range(len(df)):
            if df[[0]].iloc[i][0].isdigit() != True:
                count+=1
        df=df.iloc[count:,1:] #选取序号列为数字的行，即表格里的具体内容，不读取表头和序号列
        df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关', '处罚决定日期'] #重新对整个表格的列命名
    else:  #针对剩余和之前地区格式类似的情况
        df=df.drop(df[[7]],axis=1)#去掉备注列
        df=df.dropna(axis=0,how="all") #删除所有内容为空的行
        for i in range(len(df)):
            if df[[0]].iloc[i][0].isdigit() != True:
                count+=1
        df=df.iloc[count:,1:] #选取序号列为数字的行，即表格里的具体内容，不读取表头和序号列
        df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关', '处罚决定日期'] #重新对整个表格的列命名
    df.to_excel(path_1+today+'\\'+name+'.xlsx') #写入对应的Excel文件


#南京
def nanjing(name):
    url=path_1+today+'\\'+name+'.html'
    l=pd.read_html(url)  #用pandas安装包读取html文件，读取成功之后会在一个list里面
    df=l[9] #在list中筛选出相对应的内容
    if len(df.columns) == 6: #针对表格格式为六列的情况
        df=df.dropna(axis=0,how="any") #删掉任意内容为空的行
        df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关', '处罚决定日期'] #重新对整个表格的列命名
        df=df.iloc[1:] # 读取表格具体内容，不读取表头和序号列
        df.to_excel(path_1+today+'\\'+name+'.xlsx') #写入对应的Excel文件
    else:  #针对剩余和之前地区格式类似的情况
        count=0
        df=df.drop(df[[7]],axis=1) #去掉备注列
        df=df.dropna(axis=0,how="all") #删除所有内容为空的行
        for i in range(len(df)):
            if df[[0]].iloc[i][0].isdigit() != True:
                count+=1

        df=df.iloc[count:,1:] #选取序号列为数字的行，即表格里的具体内容，不读取表头和序号列
        df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关名称', '处罚决定日期']  #重新对整个表格的列命名
        df.to_excel(path_1+today+'\\'+name+'.xlsx') #写入对应的Excel文件



#南昌
def nanchang(name):
    url=path_1+today+'\\'+name+'.html'
    l=pd.read_html(url) #用pandas安装包读取html文件，读取成功之后会在一个list里面
    df=l[9] #在list中筛选出相对应的内容
    count=0
    df=df.drop(df[[7]],axis=1) #去掉备注列
    df=df.dropna(axis=0,how="all") #删除所有内容为空的行
    for i in range(len(df)):
        if df[[0]].iloc[i][0].isdigit() != True:
            count+=1

    df=df.iloc[count:,1:] #选取序号列为数字的行，即表格里的具体内容，不读取表头和序号列
    df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关名称', '处罚决定日期']  #重新对整个表格的列命名
    df.to_excel(path_1+today+'\\'+name+'.xlsx') #写入对应的Excel文件



#北京
def beijing(name):
    url=path_1+today+'\\'+name+'.html'
    l=pd.read_html(url) #用pandas安装包读取html文件，读取成功之后会在一个list里面
    df=l[9] #在list中筛选出相对应的内容
    count=0
    df=df.drop(df[[7]],axis=1) #去掉备注列
    df=df.dropna(axis=0,how="all") #删除所有内容为空的行
    for i in range(len(df)):
        if df[[0]].iloc[i][0].isdigit() != True:
            count+=1

    df=df.iloc[count:,1:] #选取序号列为数字的行，即表格里的具体内容，不读取表头和序号列
    df.columns=['企业名称', '处罚决定书文号', '违法行为类型', '处罚内容', '处罚决定机关名称', '处罚决定日期'] #重新对整个表格的列命名
    df.to_excel(path_1+today+'\\'+name+'.xlsx') #写入对应的Excel文件


#修改将html转换成Excel文件后的格式
def modify_excel_html(name,url):
    df = pd.read_excel(path_1+today+'\\'+name+'.xlsx')  #读取Excel文件
    nrows = df.shape[0]       #定义Excel文件里表格的总行数
    ncols = df.columns.size   #定义Excel文件里表格的列数
    workbook = xlwt.Workbook()                          #新建Excel文件
    sheet = workbook.add_sheet('Sheet1')                #新建Excel文件中的工作表并命名为'Sheet1'
    
    i = 1                     #记录在写入Excel文件时的行数
    n = 0                     #定义Excel文件里有效内容的行数

    #跳过第一行写入内容
    for row in range(nrows):
        for j in range(1, ncols):
            content = str(df.iloc[row,j])
            sheet.write(i, j-1, content)
        i += 1
    if ncols > 6:   #若表格超过6列，即存在"备注"列，则将该html所属的行政处罚链接写入"备注"列
        for i in range(1,nrows+1):
            sheet.write(i,ncols-1,url)
    else:           #若表格不存在"备注"列，则新建一列并写入该html所属的行政处罚链接
        for i in range(1,nrows+1):
            sheet.write(i,ncols,url)
    workbook.save(path_1+today+'(Excel)\\'+name+'.xls')     #保存Excel文件
    

def ab(df):
    return ''.join(df.values) #合并内容的函数

#合并全部数据
def combine_data(today):
    path = path_1+today+'(Excel)\\'

    DFs=[]

    for root, dirs, files in os.walk(path):   #第一个为起始路径，第二个为起始路径下的文件夹，第三个是起始路径下的文件
        for file in files:
            file_path=os.path.join(root,file) #将路径名和文件名组合成一个完整路
            df = pd.read_excel(file_path)
            if '序号' in df.columns:
                df=df.drop(labels='序号',axis=1) 
            if '序\n号' in df.columns:
                df=df.drop(labels='序\n号',axis=1) #针对所有表格去掉所有序号列
            if '序号.1' in df.columns:
                df=df.drop(labels='序号.1',axis=1)
            if '备注' in df.columns:
                df=df.drop(labels='备注',axis=1)
            if '备\n注' in df.columns:
                df=df.drop(labels='备\n注',axis=1) #去掉所有备注列
            if '备注.1' in df.columns:
                df=df.drop(labels='备注.1',axis=1)
            if '联系电话' in df.columns:
                df=df.drop(labels='联系电话',axis=1) #去掉联系电话列
            if '社会信用代码（组织机构代码）'in df.columns: #定位到某格式特殊的表格
                df=df.drop(labels=['社会信用代码（组织机构代码）','法人代表姓名','法人代表证件类型','法人代表证件号','报送单位名称','处罚单位地址','处罚单位电话','违法行为','违法金额','处罚执行情况','案件申请强制执行情况','预留字段/备注'],axis=1)
                df=df[['企业名称','行政处罚决定书文号','违法行为类别','处罚金额','处罚单位名称','违法行为发生时间','参考来源']] #删掉所有不必要的列，保留需要的列
                df['处罚金额']='罚款'+ df['处罚金额'] # 修改罚款金额格式为罚款多少金额
                df['违法行为发生时间']=file[0:4]+'年'+file[4:6]+'月'+file[6:8]+'日' #修改违法行为发生时间格式
            
            df=df.reset_index() 
            df=df.drop(labels='index',axis=1) #重新设置索引

            if len(df.columns)==8: #针对所有表格中表格格式为8列的情况，即金额被单独为一列的情况 
                df.columns=['企业名称','处罚决定文书号','违法行为类型','处罚内容','金额','处罚决定机关','处罚决定日期','参考来源'] #对表格的所有列重新命名
                for k in range(len(df['处罚内容'])):
                    df['处罚内容'][k]= df['处罚内容'][k] + df['金额'][k] #更新处罚内容格式为处罚内容加上金额，即罚款多少元
                df= df.drop(labels='金额',axis=1) #去掉金额列，只保留处罚内容列
            else:
                df.columns=['企业名称','处罚决定文书号','违法行为类型','处罚内容','处罚决定机关','处罚决定日期','参考来源'] #针对一般情况对表格列重命名
            for j in range(len(df['处罚决定日期'])):           
                if pd.isnull(df['处罚决定日期'][j]) == True: 
                    continue
                if len(str(df['处罚决定日期'][j]))<=7:
                    df['处罚决定日期'][j]= datetime.datetime(*xldate_as_tuple(df['处罚决定日期'][j],0)).strftime('%Y年%m月%d日')  #把一串数字变成日期       

            DFs.append(df) #把每一个表格生成的dataframe加在一个大的叫DFs的list里

    df= pd.concat(DFs,axis=0) #将生成的结果转换为一个大的dataframe
    df = df.iloc[:,:].replace('',np.nan) #将所有空格替换为nan
    df = df.dropna(axis=0,how='all') #删掉所有内容为空的行

    df=df.reset_index()
    new=df.drop(labels='index',axis=1)
    new=new.reset_index()
    new.columns=['序号','企业名称', '处罚决定文书号', '违法行为类型', '处罚内容', '处罚决定机关', '处罚决定日期','参考来源'] #重命名总表的列名
    
    #针对读取的dataframe存在截断读取内容的问题，即原文件中把一行内容分成了两行阅读，我们进行了合并操作
    for k in range(len(new['违法行为类型'])):
         if pd.isnull(new.at[k,'处罚决定日期'])==True: #如果处罚决定日期为空，即判断为截断读取内容问题的判断条件
            new['序号'][k] =  new['序号'][k-1]   #统一原本应该在同一行内容的序号

    for j in new.columns:
        new[j].fillna('', inplace=True) #将所有的nan替换为空格
    new1=new.groupby(['序号'])['企业名称'].apply(ab) 
    new2=new.groupby(['序号'])['处罚决定文书号'].apply(ab)
    new3=new.groupby(['序号'])['违法行为类型'].apply(ab)
    new4=new.groupby(['序号'])['处罚内容'].apply(ab)
    new5= new.groupby(['序号'])['处罚决定机关'].apply(ab)
    new6= new.groupby(['序号'])['处罚决定日期'].apply(ab) #生成每一列按照序号分组的合并好的结果，即按照每一列解决了截断读取的问题
    
    df0=pd.merge(new1,new2,on='序号')
    df1=pd.merge(df0,new3,on='序号')
    df2=pd.merge(df1,new4,on='序号')
    df3=pd.merge(df2,new5,on='序号')
    df4= pd.merge(df3,new6,on='序号') #将这些dataframe按照序号再重新合并起来生成总表
    df4=df4.dropna(axis=0,how='any') #去掉所有内容为空的行
    df5=new[['序号','参考来源']] 
    df6=pd.merge(df4,df5,on='序号') #将对应的参考来源合并入大表
    df6 = df6.iloc[:,:].replace('',np.nan) #将所有的空格替换为nan
    df6 = df6.fillna(method='ffill') #针对剩余的空格情况，直接按照上一行的内容填充
    df6=df6[~df6['违法行为类型'].isin([''])] 
    df6=df6.drop(labels='序号',axis=1) #去掉序号列
    
    df6=df6.reset_index() 
    df_t=df6.drop(labels='index',axis=1) #重新设置索义
    
    #循环"处罚决定日期"列，修改日期格式
    for i in range(len(df_t['处罚决定日期'])):
        date = str(df_t['处罚决定日期'][i]).replace('\n','').replace(' ','')
        if '日期' in date:    #去掉表格里可能未删除的表头
            df_t = df_t.drop(i, axis=0)
        #将不同格式的日期统一修改成格式"yyyy年mm月dd日"
        if '.' in date:
            df_t['处罚决定日期'][i] = time.strftime("%Y年%m月%d日", time.strptime(date, "%Y.%m.%d")).replace('\n','') 
        if '-' in date:
            if len(str(date)) >= 12:
                df_t['处罚决定日期'][i]= time.strftime("%Y年%m月%d日", time.strptime(date, "%Y-%m-%d00:00:00")).replace('\n','')
            else:
                date = date.replace('\xa0','')
                df_t['处罚决定日期'][i]= time.strftime("%Y年%m月%d日", time.strptime(date, "%Y-%m-%d")).replace('\n','')
        if '/' in date:
            df_t['处罚决定日期'][i]=time.strftime("%Y年%m月%d日", time.strptime(date, "%Y/%m/%d")).replace('\n','')
    df_t.columns=['企业名称','处罚决定文书号','违法行为类型','处罚内容','处罚决定机关','处罚决定日期','参考来源']

    df_t=df_t.drop_duplicates()     #删除重复行
    df_t=df_t.reset_index()
    df_t=df_t.drop(labels='index',axis=1)    
    df_t.to_excel(path_1+today+'all_data.xlsx')     #导出有全部数据的Excel文件

    df_n = df_t
    #反洗钱相关关键词筛查
    for i in range(len(df_n['违法行为类型'])):
        content = df_n['违法行为类型'][i].replace('\n','').replace(' ','')
        count = 0       #定义变量以判断该行是否有反洗钱相关的关键词
        for j in key_words:
            if j in content:    #若反洗钱相关关键词出现在"违法行为类型"列，则count加一
                count += 1
        if count == 0:          #若count值为0，即没有反洗钱相关关键词出现在"违法行为类型"列，则删除该行
            df_n=df_n.drop(i,axis=0)

    df_n=df_n.reset_index()
    df_m=df_n.drop(labels='index',axis=1)
    #添加列
    df_m['处罚金额（人民币 · 万元）'] = ''
    df_m['PBOC机构'] = ''
    df_m['年份'] = ''
    df_m['处罚对象'] = ''
    df_m['当事人'] = ''

    #提取"处罚内容"列的金额并写入"处罚金额（人民币 · 万元）"列
    for i in range(len(df_m['处罚内容'])):
        content = df_m['处罚内容'][i]    #定义该行"处罚内容"列的值
        if len(content) >= 20:          #若该行"处罚内容"列的值过长，则需后续人工填写
            df_m['处罚金额（人民币 · 万元）'][i] = ''
        else:
            money = re.findall('([\d]+[,.]*[\d]*)',content) #提取该行"处罚内容"列的值里的数值
            if money == []:
                df_m['处罚金额（人民币 · 万元）'][i] = ''      #若未提取出数值，则需后续人工填写
            else:
                money = money[0].replace(',','')    #修改提取出的数值格式
                if '万' in content:
                    money = '%.2f' % float(money)   #若该行"处罚内容"列的值存在"万"字，则保留两位小数并写入"处罚金额（人民币 · 万元）"列
                    df_m['处罚金额（人民币 · 万元）'][i] = float(money)
                else:
                    amount = float(money)/10000     #若该行"处罚内容"列的值不存在"万"字，则将数值除以10000并保留两位小数
                    amount = '%.2f' % amount
                    df_m['处罚金额（人民币 · 万元）'][i] = float(amount)   #写入"处罚金额（人民币 · 万元）"列
                    
    #提取地名拼音并查找表格df_inf里地名拼音对应的支行名称，写入"PBOC机构"列
    for i in range(len(df_m['参考来源'])):
        place = df_m['参考来源'][i].split('/')[2].split('.')[0]        #提取该行所属支行所在城市的拼音
        content = str(df_inf.loc[df_inf['place']==place,'name'])      #通过拼音查找表格df_inf里对应的支行名称
        df_m['PBOC机构'][i] = ''.join(re.findall(findCHN, content))   #修改格式，写入"PBOC机构"列
        
    #提取"处罚决定日期"列的年份并写入"年份"列
    for i in range(len(df_m['处罚决定日期'])):
        content = df_m['处罚决定日期'][i].replace(' ','').replace('\n','')    #提取年份
        df_m['年份'][i] = content[0:4]                                       #写入"年份"列
        
    #对"企业名称"列进行判断该行为个人或机构并写入"处罚对象"列与"当事人"列
    for i in range(len(df_m['企业名称'])):
        content = df_m['企业名称'][i].replace('\n','').replace(' ','').replace('（','(').replace('）',')')
        if ('(' in content and ')' in content):     #若该行"企业名称"列存在括号，则进行条件判断
            content_p = content.split('(') 
            content_p = content_p[1].split(')')
            content_p = content_p[0]                #提取括号内的文字
            if len(content_p) > 5:                  
                df_m['处罚对象'][i] = '个人'          #若括号内的文字内容较长，则定义该行为"个人"
                df_m['当事人'][i]=content            #将该行"企业名称"列的内容写入"当事人"列
                df_m['企业名称'][i]=''               #将该行"企业名称"列赋值为""以方便后续修改
                continue
            else:
                if '公司' not in content:            #若括号内文字内容较短，则判断内容是否存在"公司"
                    df_m['处罚对象'][i] = '个人'      #若不存在"公司"，则定义该行为"个人"
                    df_m['当事人'][i]=content
                    df_m['企业名称'][i]=''
                    continue
        elif len(content) <= 5:             #若该行"企业名称"列内容短于5，则定义该行为"个人"
            df_m['处罚对象'][i] = '个人'
            df_m['当事人'][i]=content
            df_m['企业名称'][i]=''
            continue
        elif '责任人' in content:            #若该行"企业名称"列存在"负责人"，则定义该行为"个人"
            df_m['处罚对象'][i] = '个人'
            df_m['当事人'][i]=content
            df_m['企业名称'][i]=''
            continue
        df_m['处罚对象'][i] = '机构'          #若均不符合上述情况，则定义该行为"机构"
        df_m['当事人'][i]='nil'              #将"当事人"列赋值为"nil"

    df_m['企业名称']=df_m.iloc[:,0].replace('',np.nan)  #将"企业名称"列所有为""的值替换为"NaN"
    df_m = df_m.fillna(method='ffill')                 #将整张表格所有为"NaN"的值替换为该列上一个不为"NaN"的值

    #添加省份和区域列,参照Kary和Dora的advice
    city={'上海分行':'上海','天津分行':'天津','沈阳分行':'辽宁','南京分行':'江苏','济南分行':'山东','武汉分行':'湖北',
    '广州分行':'广东','成都分行':'四川','西安分行':'陕西','北京营业管理部':'北京','重庆营业管理部':'重庆',
    '石家庄中心支行':'河北','太原中心支行':'山西','呼和浩特中心支行':'内蒙古','长春中心支行':'吉林',
     '哈尔滨中心支行':'黑龙江','杭州中心支行':'浙江','福州中心支行':'福建','合肥中心支行':'安徽',
    '郑州中心支行':'河南','长沙中心支行':'湖南','南昌中心支行':'江西','南宁中心支行':'广西','海口中心支行':'海南',
    '昆明中心支行':'云南','贵阳中心支行':'贵州','拉萨中心支行':'西藏','兰州中心支行':'甘肃','西宁中心支行':'青海',
    '银川中心支行':'宁夏','乌鲁木齐中心支行':'新疆','深圳市中心支行':'广东','大连市中心':'辽宁','青岛市中心支行':'山东',
     '宁波市中心支行':'浙江','厦门市中心支行':'福建'}
    df_m['省份']=''
    df_m['区域']=''
    for i in range(len(df_m['PBOC机构'])):
        df_m['省份'][i]=city[df_m['PBOC机构'][i]]
    Area={'山东':'华东','江苏':'华东','安徽':'华东','浙江':'华东','福建':'华东','江西':'华东','上海':'华东',
        '广东':'华南', '广西':'华南','海南':'华南', '河北':'华北','山西':'华北','北京':'华北','天津':'华北','内蒙古':'华北',
        '湖北':'华中','湖南':'华中','河南':'华中','辽宁':'东北','吉林':'东北','黑龙江':'东北',
         '四川':'西南', '云南':'西南', '贵州':'西南', '重庆':'西南','西藏':'西南',
         '陕西':'西北','宁夏':'西北','新疆':'西北','青海':'西北','甘肃':'西北'}
    for j in range(len(df_m['省份'])):
        df_m['区域'][j]=Area[df_m['省份'][j]]
        df_m=df_m[['Unnamed: 0','PBOC机构','省份','区域','处罚决定机关','年份','处罚决定日期','处罚决定文书号','企业名称','处罚对象','当事人','处罚内容','处罚金额（人民币 · 万元）',
                    '违法行为类型','参考来源']]
        df_m=df_m.drop(labels='Unnamed: 0',axis=1)
        
    #添加企业类型和子类型列，大部分企业名称可以成功分类，少部分需要人工审阅，参照Olive的advice
    df_m['企业类型']=''
    df_m['企业子类型']=''
    l=['招商银行','广发银行','光大银行','平安银行','浦东发展银行','中信银行','华夏银行','民生银行','兴业银行','恒丰银行','浙商银行','渤海银行']
    for i in range(len(df_m['企业名称'])):
        df_m['企业名称'][i]=df_m['企业名称'][i].replace('\n','')
        if '银行' in df_m['企业名称'][i] or '农村信用' in df_m['企业名称'][i]:
            df_m['企业类型'][i]='银行'
            if '农村商业银行股份'in df_m['企业名称'][i] or '农村银行股份'in df_m['企业名称'][i] or '农村商业银行'in df_m['企业名称'][i] or '农村合作银行' in df_m['企业名称'][i]:
                df_m['企业子类型'][i] = '农村商业银行'
            elif '银行股份' in df_m['企业名称'][i]:
                df_m['企业子类型'][i] = '股份制商业银行'
            elif '村镇银行'in df_m['企业名称'][i]:
                df_m['企业子类型'][i] = '村镇银行'
            elif '农村信用社' in df_m['企业名称'][i] or '联社' in df_m['企业名称'][i]:
                df_m['企业子类型'][i] = '农村信用社'
            elif '工商' in df_m['企业名称'][i] or '中国农业银行'in df_m['企业名称'][i] or '中国银行' in df_m['企业名称'][i] or '建设银行' in df_m['企业名称'][i] or '交通银行' in df_m['企业名称'][i] or '中国邮政储蓄银行' in df_m['企业名称'][i]:
                df_m['企业子类型'][i] = '国有商业银行'
            elif '国家开发银行' in df_m['企业名称'][i] or '农业发展银行' in df_m['企业名称'][i] or '进出口银行' in df_m['企业名称'][i]:
                df_m['企业子类型'][i] = '政策性银行'
            else:
                for j in l:
                    if j in df_m['企业名称'][i]:
                        df_m['企业子类型'][i] = '股份制商业银行'
        elif '证券' in df_m['企业名称'][i]:
            df_m['企业类型'][i]='证券'
        elif '期货' in df_m['企业名称'][i]:
            df_m['企业类型'][i]='期货'
        elif '基金' in df_m['企业名称'][i]:
            df_m['企业类型'][i]='基金'
        elif '保险' in df_m['企业名称'][i]:
            df_m['企业类型'][i]='保险'
        elif '信托' in df_m['企业名称'][i]:
            df_m['企业类型'][i]='信托'
        elif '金融控股' in df_m['企业名称'][i]:
            df_m['企业类型'][i]='金融资管'
        elif '支付' in df_m['企业名称'][i] or '科技' in df_m['企业名称'][i] or '商务' in df_m['企业名称'][i]:
            df_m['企业类型'][i]='支付机构'
    df_m=df_m[['PBOC机构','省份','区域','处罚决定机关','年份','处罚决定日期','处罚决定文书号','企业名称','企业类型','企业子类型','处罚对象','当事人','处罚内容','处罚金额（人民币 · 万元）',
                '违法行为类型','参考来源']]

    df_m=df_m.drop_duplicates()     #删除重复行
    df_m=df_m.reset_index()
    df_m=df_m.drop(labels='index',axis=1)
    
    
    df_m.to_excel(path_1+today+'.xlsx')


if __name__ == "__main__":
     main()
     print("爬取完毕！")


